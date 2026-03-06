
"""
json_to_docx_image.py (v4)

Read SLAC parsed JSON files (structure tree) and reconstruct:
1) A Word document (.docx) listing units in original order, with a compact metadata header
   before each unit, including parent info (parent_id + parent preview).
2) A tree-structure image (.png) for each JSON with clearer layout:
   - nodes are circles with unit_id inside
   - image title includes file name (stem) and doc_name
   - default renders a "summary tree" (root + title/heading nodes + ancestors) for readability
   - optional full tree rendering

Default paths (can be overridden by CLI):
- input_json_dir:  D:\\code\\Github\\SLAC-test\\SLAC\\data\\parsed_json
- output_docx_dir: D:\\code\\Github\\SLAC-test\\SLAC\\data\\parsed_word
- output_img_dir:  D:\\code\\Github\\SLAC-test\\SLAC\\data\\structure_image
- error_log_dir:   D:\\code\\Github\\SLAC-test\\SLAC\\data\\structure_image\\_errors

Dependencies:
- python-docx
- matplotlib
(Optional for better graphs)
- graphviz (python package) + Graphviz "dot" installed in OS PATH

Run:
python json_to_docx_image.py
or
python json_to_docx_image.py --input_json_dir ... --output_docx_dir ... --output_img_dir ...
"""

from __future__ import annotations

import os
import json
import argparse
import traceback
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Set

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt


# ----------------------------
# Helpers
# ----------------------------
def safe_str(x: Any) -> str:
    if x is None:
        return ""
    try:
        return str(x)
    except Exception:
        return repr(x)

def ensure_dir(p: Path) -> None:
    p.mkdir(parents=True, exist_ok=True)

def iter_json_files(root: Path) -> List[Path]:
    if not root.exists():
        return []
    return sorted([p for p in root.rglob("*.json") if p.is_file()])

def load_json(p: Path) -> Dict[str, Any]:
    with open(p, "r", encoding="utf-8") as f:
        return json.load(f)

def dump_error_log(log_path: Path, msg: str) -> None:
    ensure_dir(log_path.parent)
    with open(log_path, "a", encoding="utf-8") as f:
        f.write(msg.rstrip() + "\n")

def normalize_spaces(s: str) -> str:
    s = (s or "").replace("\u00a0", " ")
    s = " ".join(s.split())
    return s.strip()

def has_cjk(s: str) -> bool:
    import re
    return bool(re.search(r"[\u4e00-\u9fff]", s or ""))

def sentence_end_punct(s: str) -> bool:
    import re
    s = (s or "").rstrip()
    return bool(re.search(r"[.!?;:。！？；：]$", s))

def trunc(s: str, n: int) -> str:
    s = normalize_spaces(s)
    return s if len(s) <= n else (s[: max(0, n - 1)] + "…")


# ----------------------------
# DOCX writing
# ----------------------------
def set_run_font(run, zh_font: str, en_font: str, size_pt: int = 11, color: Optional[RGBColor] = None, mono: bool = False):
    run.font.size = Pt(size_pt)
    # latin font
    run.font.name = en_font if not mono else "Consolas"
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), zh_font)
    except Exception:
        pass
    if color is not None:
        run.font.color.rgb = color

def add_doc_header(doc: Document, doc_id: str, doc_name: str, language: str, source_json: Path, units_count: int,
                   zh_font: str, en_font: str) -> None:
    doc.add_heading(doc_name or "unknown_title", level=0)
    p = doc.add_paragraph()
    r = p.add_run(f"doc_id: {doc_id}    language: {language}    units: {units_count}")
    set_run_font(r, zh_font, en_font, size_pt=10, color=RGBColor(70, 70, 70))
    p2 = doc.add_paragraph()
    r2 = p2.add_run(f"source_json: {str(source_json)}")
    set_run_font(r2, zh_font, en_font, size_pt=9, color=RGBColor(90, 90, 90))
    doc.add_paragraph("")

def normalize_text_for_docx(text: str) -> List[str]:
    """
    Convert raw unit text to a list of readable paragraphs.

    Rules:
    - collapse Windows newlines
    - split into paragraph blocks by blank lines
    - within a block, join lines:
        * for zh: join without extra spaces
        * for en: join with spaces, and fix hyphenation like "inter-\nnational" -> "international"
    - drop empty paragraphs
    """
    if not text:
        return []
    t = text.replace("\r\n", "\n").replace("\r", "\n")
    # strip trailing spaces per line
    lines = [ln.rstrip() for ln in t.split("\n")]
    # collapse excessive blank lines to single blank delimiter
    blocks: List[List[str]] = []
    cur: List[str] = []
    for ln in lines:
        if not ln.strip():
            if cur:
                blocks.append(cur)
                cur = []
            continue
        cur.append(ln)
    if cur:
        blocks.append(cur)

    out_paras: List[str] = []
    for blk in blocks:
        # remove decorative-ish single lines that are just symbols
        blk2 = [x for x in blk if normalize_spaces(x)]
        if not blk2:
            continue
        zh = any(has_cjk(x) for x in blk2)
        if zh:
            # join, but keep spaces that exist in original when it looks like English tokens
            joined = "".join([x.strip() for x in blk2])
            out_paras.append(normalize_spaces(joined))
        else:
            merged_parts: List[str] = []
            for i, ln in enumerate(blk2):
                s = ln.strip()
                if not s:
                    continue
                if merged_parts:
                    prev = merged_parts[-1]
                    # hyphenation fix: "inter-" + "national" => "international" if next starts with lowercase
                    if prev.endswith("-") and s and s[0].islower():
                        merged_parts[-1] = prev[:-1] + s
                    else:
                        # if prev ends with sentence punctuation, keep a space anyway; otherwise space
                        merged_parts.append(s)
                else:
                    merged_parts.append(s)
            joined = " ".join(merged_parts)
            out_paras.append(normalize_spaces(joined))

    # Final cleanup: drop very short empty-ish
    out_paras = [p for p in out_paras if p and p.strip()]
    return out_paras

def add_meta_paragraph(doc: Document, meta: str, zh_font: str, en_font: str, level: int = 0) -> None:
    p = doc.add_paragraph()
    # indent with level but cap
    indent_in = min(1.5, 0.25 * max(0, int(level)))
    if indent_in > 0:
        p.paragraph_format.left_indent = Inches(indent_in)
    r = p.add_run(meta)
    set_run_font(r, zh_font, en_font, size_pt=9, color=RGBColor(80, 80, 80), mono=True)

def add_unit_body(doc: Document, paras: List[str], zh_font: str, en_font: str, level: int, utype: str) -> None:
    """
    Write unit text as a small set of paragraphs with indentation.
    """
    indent_in = min(1.5, 0.25 * max(0, int(level)))
    for j, para in enumerate(paras):
        if not para.strip():
            continue
        # For headings/titles: use built-in heading for the first paragraph only
        if j == 0 and utype in {"title", "heading"}:
            # Heading levels in Word: 1..9, with 0 reserved for doc title
            h_level = min(9, max(1, int(level) + 1))
            doc.add_heading(para, level=h_level)
            continue

        p = doc.add_paragraph()
        if indent_in > 0:
            p.paragraph_format.left_indent = Inches(indent_in)
        r = p.add_run(para)
        set_run_font(r, zh_font, en_font, size_pt=11)

def json_to_docx(doc_json: Dict[str, Any], out_docx: Path, source_json: Path,
                 zh_font: str, en_font: str, add_parent_preview: bool = True) -> None:
    ensure_dir(out_docx.parent)
    doc = Document()

    doc_id = safe_str(doc_json.get("doc_id", source_json.stem))
    doc_name = safe_str(doc_json.get("doc_name", "unknown_title"))
    language = safe_str(doc_json.get("language", "other"))
    units = doc_json.get("units", []) or []
    units_count = len(units)

    # build unit lookup
    unit_by_id: Dict[int, Dict[str, Any]] = {}
    for u in units:
        try:
            unit_by_id[int(u.get("unit_id"))] = u
        except Exception:
            continue

    add_doc_header(doc, doc_id, doc_name, language, source_json, units_count, zh_font, en_font)

    for u in units:
        try:
            uid = int(u.get("unit_id"))
        except Exception:
            continue
        level = u.get("level", 0)
        try:
            level_i = int(level)
        except Exception:
            level_i = 0

        pid = u.get("parent_id", None)
        pid_i: Optional[int] = None
        if pid is not None:
            try:
                pid_i = int(pid)
            except Exception:
                pid_i = None

        utype = safe_str(u.get("type", ""))
        np = u.get("num_prefix", None)
        mk = u.get("marker_type", None)
        h = u.get("unit_hash", None)

        parent_desc = ""
        if add_parent_preview:
            if pid_i is None:
                parent_desc = "parent=None"
            else:
                pu = unit_by_id.get(pid_i)
                if pu is None:
                    parent_desc = f"parent={pid_i}"
                else:
                    p_np = pu.get("num_prefix") or ""
                    p_type = pu.get("type") or ""
                    p_preview = trunc(safe_str(pu.get("text", "")), 60)
                    # avoid giant meta strings
                    parent_desc = f"parent={pid_i}({p_type}{(':'+p_np) if p_np else ''}) {p_preview}"

        meta = (
            f"[unit_id={uid} | level={level_i} | parent_id={pid_i} | type={utype}"
            f"{' | num_prefix=' + safe_str(np) if np is not None else ''}"
            f"{' | marker_type=' + safe_str(mk) if mk is not None else ''}"
            f"{' | hash=' + safe_str(h) if h is not None else ''}"
            f"{' | ' + parent_desc if parent_desc else ''}"
            f"]"
        )
        add_meta_paragraph(doc, meta, zh_font, en_font, level=level_i)

        paras = normalize_text_for_docx(safe_str(u.get("text", "")))
        if not paras:
            paras = ["(empty)"]
        add_unit_body(doc, paras, zh_font, en_font, level=level_i, utype=utype)

        # spacer: one blank paragraph only
        doc.add_paragraph("")

    doc.save(out_docx)


# ----------------------------
# Tree rendering
# ----------------------------
def build_parent_map(units: List[Dict[str, Any]]) -> Tuple[int, Dict[int, Optional[int]], Dict[int, Dict[str, Any]]]:
    root_id = 0
    parent: Dict[int, Optional[int]] = {}
    by_id: Dict[int, Dict[str, Any]] = {}
    for u in units:
        try:
            uid = int(u.get("unit_id"))
        except Exception:
            continue
        by_id[uid] = u
        pid = u.get("parent_id", None)
        if pid is None:
            root_id = uid
            parent[uid] = None
        else:
            try:
                parent[uid] = int(pid)
            except Exception:
                parent[uid] = None
    if 0 in by_id:
        root_id = 0
        parent[0] = None
    return root_id, parent, by_id

def build_children_map(parent: Dict[int, Optional[int]]) -> Dict[int, List[int]]:
    children: Dict[int, List[int]] = {}
    for uid, pid in parent.items():
        if pid is None:
            continue
        children.setdefault(pid, []).append(uid)
    for k in list(children.keys()):
        children[k] = sorted(children[k])
    return children

def compute_levels_from_units(units: List[Dict[str, Any]]) -> Dict[int, int]:
    lvl: Dict[int, int] = {}
    for u in units:
        try:
            uid = int(u.get("unit_id"))
            level = int(u.get("level", 0))
            lvl[uid] = level
        except Exception:
            continue
    return lvl

def select_summary_nodes(root: int, parent: Dict[int, Optional[int]], by_id: Dict[int, Dict[str, Any]],
                         keep_types: Set[str], max_depth: int) -> Set[int]:
    """
    Keep root + nodes whose type in keep_types and level <= max_depth, plus their ancestors.
    """
    keep: Set[int] = set([root])
    for uid, u in by_id.items():
        try:
            lvl = int(u.get("level", 0))
        except Exception:
            lvl = 0
        if lvl > max_depth:
            continue
        t = (u.get("type") or "").strip()
        if t in keep_types:
            keep.add(uid)
            # add ancestors
            cur = uid
            steps = 0
            while True:
                pid = parent.get(cur)
                if pid is None:
                    break
                keep.add(pid)
                cur = pid
                steps += 1
                if steps > 20000:
                    break
    return keep

def tree_layout_tidy(root: int, children: Dict[int, List[int]], levels: Dict[int, int], nodes: Set[int], y_gap: float = 1.6) -> Dict[int, Tuple[float, float]]:
    """
    Simple tidy layout:
      - x assigned by DFS leaf order within selected nodes
      - internal node x = mean(children x)
      - y = -level * y_gap
    """
    pos: Dict[int, Tuple[float, float]] = {}
    next_x = 0.0

    def dfs(u: int) -> float:
        nonlocal next_x
        ch_all = children.get(u, [])
        ch = [v for v in ch_all if v in nodes]
        if not ch:
            x = next_x
            next_x += 1.0
            y = -float(levels.get(u, 0)) * float(y_gap)
            pos[u] = (x, y)
            return x
        xs = [dfs(v) for v in ch]
        x = sum(xs) / len(xs) if xs else next_x
        y = -float(levels.get(u, 0)) * float(y_gap)
        pos[u] = (x, y)
        return x

    if root in nodes:
        dfs(root)

    # place remaining nodes (disconnected) to the right
    for uid in sorted(nodes):
        if uid not in pos:
            x = next_x
            next_x += 1.0
            y = -float(levels.get(uid, 0)) * float(y_gap)
            pos[uid] = (x, y)
    return pos

def draw_tree_matplotlib(units: List[Dict[str, Any]], out_png: Path, title: str,
                         mode: str = "summary",
                         keep_types: Optional[Set[str]] = None,
                         max_depth: int = 8,
                         max_nodes: int = 800,
                         dpi: int = 220,
                         y_gap: float = 1.6) -> None:
    ensure_dir(out_png.parent)
    root, parent, by_id = build_parent_map(units)
    children = build_children_map(parent)
    levels = compute_levels_from_units(units)

    if keep_types is None:
        keep_types = {"title", "heading"}

    if mode == "full":
        nodes = set(by_id.keys())
    else:
        nodes = select_summary_nodes(root, parent, by_id, keep_types=keep_types, max_depth=max_depth)

    # if too large, shrink by depth first, then by types only
    if len(nodes) > max_nodes and mode != "full":
        # tighten depth
        nodes = set([n for n in nodes if levels.get(n, 0) <= max(2, max_depth - 2)])
        nodes.add(root)
    if len(nodes) > max_nodes:
        # last resort: keep only root + keep_types
        nodes = set([root])
        for uid, u in by_id.items():
            if (u.get("type") or "").strip() in keep_types and levels.get(uid, 0) <= max_depth:
                nodes.add(uid)

    pos = tree_layout_tidy(root, children, levels, nodes, y_gap=y_gap)

    # figure sizing: scale with node count and depth
    max_level = max([levels.get(n, 0) for n in nodes], default=0)
    leaf_count = sum(1 for n in nodes if len([c for c in children.get(n, []) if c in nodes]) == 0)
    width = max(10.0, min(80.0, 0.35 * max(10, leaf_count)))
    height = max(6.0, min(90.0, 1.2 * max(6, float(y_gap) * (max_level + 3))))

    fig = plt.figure(figsize=(width, height), dpi=dpi)
    ax = fig.add_subplot(111)
    ax.set_axis_off()

    # draw edges
    for uid in nodes:
        pid = parent.get(uid)
        if pid is None or pid not in nodes:
            continue
        x1, y1 = pos[pid]
        x2, y2 = pos[uid]
        ax.plot([x1, x2], [y1, y2], linewidth=0.8)

    # draw nodes as circles with unit_id
    xs = [pos[n][0] for n in nodes]
    ys = [pos[n][1] for n in nodes]
    # node size scales: summary bigger, full smaller
    base_s = 380 if mode != "full" else 120
    ax.scatter(xs, ys, s=base_s, marker="o", linewidths=0.8, edgecolors="black", facecolors="white", zorder=3)

    # labels inside circles
    # font size depends on node count
    if len(nodes) <= 120:
        fs = 10
    elif len(nodes) <= 300:
        fs = 8
    else:
        fs = 6

    for n in nodes:
        x, y = pos[n]
        ax.text(x, y, str(n), ha="center", va="center", fontsize=fs, zorder=4)

    fig.suptitle(title, fontsize=12, y=0.995)
    plt.tight_layout(rect=[0, 0, 1, 0.985])
    fig.savefig(out_png, bbox_inches="tight")
    plt.close(fig)

def try_draw_tree_graphviz(units: List[Dict[str, Any]], out_png: Path, title: str,
                           mode: str = "summary",
                           keep_types: Optional[Set[str]] = None,
                           max_depth: int = 8,
                           max_nodes: int = 1200,
                           nodesep: float = 0.35,
                           ranksep: float = 0.85,
                           dpi: int = 220) -> bool:
    """
    Render via Graphviz if available. Returns True if succeeded.
    """
    try:
        from graphviz import Digraph
    except Exception:
        return False

    root, parent, by_id = build_parent_map(units)
    children = build_children_map(parent)
    levels = compute_levels_from_units(units)

    if keep_types is None:
        keep_types = {"title", "heading"}

    if mode == "full":
        nodes = set(by_id.keys())
    else:
        nodes = select_summary_nodes(root, parent, by_id, keep_types=keep_types, max_depth=max_depth)

    if len(nodes) > max_nodes and mode != "full":
        nodes = set([n for n in nodes if levels.get(n, 0) <= max(2, max_depth - 2)])
        nodes.add(root)
    if len(nodes) > max_nodes:
        # too big for graphviz export safely
        return False

    ensure_dir(out_png.parent)

    dot = Digraph("tree", format="png")
    dot.attr(dpi=str(dpi))
    dot.attr(rankdir="TB", splines="line", nodesep=str(nodesep), ranksep=str(ranksep))
    dot.attr(label=title, labelloc="t", fontsize="18")

    # node style: circle with unit_id label
    dot.attr("node", shape="circle", fixedsize="true", width="0.45", fontsize="10", style="filled", fillcolor="white")

    for n in sorted(nodes):
        dot.node(str(n), label=str(n))

    for uid in sorted(nodes):
        pid = parent.get(uid)
        if pid is None or pid not in nodes:
            continue
        dot.edge(str(pid), str(uid))

    try:
        dot.render(str(out_png.with_suffix("")), cleanup=True)
        # graphviz renders to out_png.with_suffix(".png")
        return True
    except Exception:
        return False

def draw_tree_image(units: List[Dict[str, Any]], out_png: Path, title: str,
                    tree_mode: str = "summary",
                    max_depth: int = 8,
                    max_nodes: int = 800,
                    dpi: int = 220,
                    prefer_graphviz: bool = True,
                    y_gap: float = 1.6,
                    nodesep: float = 0.35,
                    ranksep: float = 0.85) -> None:
    keep_types = {"title", "heading"}
    ok = False
    if prefer_graphviz:
        ok = try_draw_tree_graphviz(
            units, out_png, title,
            mode=("full" if tree_mode == "full" else "summary"),
            keep_types=keep_types,
            max_depth=max_depth,
            max_nodes=max(1200, max_nodes),
            nodesep=nodesep,
            ranksep=ranksep,
            dpi=dpi,
        )
    if not ok:
        draw_tree_matplotlib(
            units, out_png, title,
            mode=("full" if tree_mode == "full" else "summary"),
            keep_types=keep_types,
            max_depth=max_depth,
            max_nodes=max_nodes,
            dpi=dpi,
            y_gap=y_gap,
        )


# ----------------------------
# Batch runner
# ----------------------------
def process_one(json_path: Path,
                out_docx_dir: Path,
                out_img_dir: Path,
                err_log_dir: Path,
                tree_mode: str,
                max_depth: int,
                max_nodes: int,
                dpi: int,
                prefer_graphviz: bool,
                y_gap: float,
                nodesep: float,
                ranksep: float,
                zh_font: str,
                en_font: str) -> None:
    stem = json_path.stem
    out_docx = out_docx_dir / f"{stem}.docx"
    out_png = out_img_dir / f"{stem}.png"
    err_log = err_log_dir / f"{stem}.err.log"

    try:
        data = load_json(json_path)
        units = data.get("units", []) or []

        # DOCX
        json_to_docx(data, out_docx, json_path, zh_font=zh_font, en_font=en_font, add_parent_preview=True)

        # Image
        doc_name = safe_str(data.get("doc_name", "unknown_title"))
        title = f"{stem} | {doc_name}" if doc_name else stem
        draw_tree_image(units, out_png, title=title,
                        tree_mode=tree_mode,
                        max_depth=max_depth,
                        max_nodes=max_nodes,
                        dpi=dpi,
                        prefer_graphviz=prefer_graphviz)

    except Exception:
        dump_error_log(err_log, f"[ERROR] {json_path}\n{traceback.format_exc()}\n")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--input_json_dir", default=r"D:\code\Github\SLAC-test\SLAC\data\parsed_json",
                    help="Directory containing parsed JSON files (recursive).")
    ap.add_argument("--output_docx_dir", default=r"D:\code\Github\SLAC-test\SLAC\data\parsed_word",
                    help="Output directory for reconstructed .docx files.")
    ap.add_argument("--output_img_dir", default=r"D:\code\Github\SLAC-test\SLAC\data\structure_image",
                    help="Output directory for tree .png images.")
    ap.add_argument("--error_log_dir", default=r"D:\code\Github\SLAC-test\SLAC\data\structure_image\_errors",
                    help="Directory to write per-json error logs.")

    # Word fonts
    ap.add_argument("--zh_font", default="宋体", help="Word Chinese font name (East Asia).")
    ap.add_argument("--en_font", default="Calibri", help="Word Latin font name.")

    # Tree rendering
    ap.add_argument("--tree_mode", choices=["summary", "full"], default="summary",
                    help="Tree image mode: summary (recommended) or full (may be unreadable for huge trees).")
    ap.add_argument("--max_depth", type=int, default=8, help="Max depth for summary tree.")
    ap.add_argument("--max_nodes", type=int, default=800, help="Max nodes to draw (summary will be pruned if exceeded).")
    ap.add_argument("--dpi", type=int, default=220, help="PNG DPI.")
    ap.add_argument("--y_gap", type=float, default=1.6,
                    help="Vertical spacing multiplier between tree levels (matplotlib fallback).")
    ap.add_argument("--nodesep", type=float, default=0.35,
                    help="Graphviz horizontal node separation (dot).")
    ap.add_argument("--ranksep", type=float, default=0.85,
                    help="Graphviz vertical rank separation between levels (dot).")
    ap.add_argument("--prefer_graphviz", action="store_true", help="Prefer Graphviz dot rendering if available.")
    ap.add_argument("--no_prefer_graphviz", action="store_true", help="Disable Graphviz; use matplotlib only.")

    args = ap.parse_args()

    prefer_graphviz = bool(args.prefer_graphviz) and (not bool(args.no_prefer_graphviz))
    # If neither flag is given, default to True (auto) because it falls back safely.
    if (not args.prefer_graphviz) and (not args.no_prefer_graphviz):
        prefer_graphviz = True

    in_dir = Path(args.input_json_dir)
    out_docx_dir = Path(args.output_docx_dir)
    out_img_dir = Path(args.output_img_dir)
    err_log_dir = Path(args.error_log_dir)

    ensure_dir(out_docx_dir)
    ensure_dir(out_img_dir)
    ensure_dir(err_log_dir)

    files = iter_json_files(in_dir)
    if not files:
        print(f"[INFO] No json found under: {in_dir}")
        return

    print(f"[INFO] Found {len(files)} json files. Start processing...")
    ok = 0
    for i, jp in enumerate(files, 1):
        print(f"[{i}/{len(files)}] {jp}")
        try:
            process_one(
                jp,
                out_docx_dir=out_docx_dir,
                out_img_dir=out_img_dir,
                err_log_dir=err_log_dir,
                tree_mode=args.tree_mode,
                max_depth=args.max_depth,
                max_nodes=args.max_nodes,
                dpi=args.dpi,
                prefer_graphviz=prefer_graphviz,
                y_gap=args.y_gap,
                nodesep=args.nodesep,
                ranksep=args.ranksep,
                zh_font=args.zh_font,
                en_font=args.en_font,
            )
            ok += 1
        except Exception:
            # process_one already logs, but keep batch alive
            pass

    print(f"[DONE] Processed: {ok}/{len(files)}")
    print(f"[DONE] DOCX dir: {out_docx_dir}")
    print(f"[DONE] IMG dir:  {out_img_dir}")
    print(f"[DONE] Errors:   {err_log_dir}")


if __name__ == "__main__":
    main()
