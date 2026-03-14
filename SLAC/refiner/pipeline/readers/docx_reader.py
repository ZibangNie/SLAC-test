"""
DOCX reader for refiner pipeline.

Main responsibility:
- read docx
- extract paragraph text
- provide unified document object
"""
from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, Iterator, List, Optional, Sequence, Union

from ..utils.ids import stable_doc_id
from ..utils.io import iter_input_files


PathLike = Union[str, Path]


def _import_python_docx():
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        raise ImportError(
            "python-docx is required for docx_reader.py. "
            "Install it with: pip install python-docx"
        ) from e
    return Document


def _nonempty_lines(lines: List[str]) -> List[str]:
    out: List[str] = []
    for line in lines:
        x = " ".join((line or "").split()).strip()
        if x:
            out.append(x)
    return out


def _extract_docx_text(docx_path: Path) -> Dict[str, Any]:
    Document = _import_python_docx()
    doc = Document(str(docx_path))

    paragraph_lines = _nonempty_lines([p.text for p in doc.paragraphs])

    table_lines: List[str] = []
    for tb in doc.tables:
        for row in tb.rows:
            cells = []
            for cell in row.cells:
                cell_text = " ".join((cell.text or "").split()).strip()
                if cell_text:
                    cells.append(cell_text)
            if cells:
                table_lines.append(" | ".join(cells))

    all_lines: List[str] = []
    if paragraph_lines:
        all_lines.extend(paragraph_lines)
    if table_lines:
        if all_lines:
            all_lines.append("")
        all_lines.extend(table_lines)

    raw_text = "\n".join(all_lines).strip()

    title_hint = paragraph_lines[0] if paragraph_lines else None

    return {
        "raw_text": raw_text,
        "title_hint": title_hint,
        "num_paragraphs": len(doc.paragraphs),
        "num_nonempty_paragraphs": len(paragraph_lines),
        "num_tables": len(doc.tables),
        "num_table_rows_kept": len(table_lines),
    }


def read_docx_document(
    path: PathLike,
    *,
    doc_id: Optional[str] = None,
    extra_meta: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    docx_path = Path(path).expanduser().resolve()
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file does not exist: {docx_path}")
    if docx_path.suffix.lower() != ".docx":
        raise ValueError(f"Not a DOCX file: {docx_path}")

    info = _extract_docx_text(docx_path)
    final_doc_id = stable_doc_id(
        explicit_doc_id=doc_id,
        source_path=docx_path,
        title_hint=info.get("title_hint"),
    )

    meta: Dict[str, Any] = {
        "reader": "docx_reader",
        "num_chars": len(info["raw_text"]),
        "title_hint": info.get("title_hint"),
        "num_paragraphs": info["num_paragraphs"],
        "num_nonempty_paragraphs": info["num_nonempty_paragraphs"],
        "num_tables": info["num_tables"],
        "num_table_rows_kept": info["num_table_rows_kept"],
    }
    if extra_meta:
        meta.update(extra_meta)

    return {
        "doc_id": final_doc_id,
        "source_path": docx_path.as_posix(),
        "source_type": "docx",
        "raw_text": info["raw_text"],
        "meta": meta,
    }


def iter_docx_documents(
    input_paths: Sequence[PathLike],
    *,
    recursive: bool = True,
) -> Iterator[Dict[str, Any]]:
    files = iter_input_files(
        input_paths,
        suffixes=[".docx"],
        recursive=recursive,
        sort_result=True,
        allow_hidden=False,
    )
    for fp in files:
        yield read_docx_document(fp)