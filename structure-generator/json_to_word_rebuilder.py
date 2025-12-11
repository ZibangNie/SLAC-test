"""
tree_to_word_batch.py

功能：
  从结构树 JSON（*.tree.json）批量生成标注好的 Word 文档和结构树图片：
  - 每个 unit 前有一行 meta 信息：[unit_id=..., type=..., level=..., parent_id=...]
  - 下方一行是对应文本，按 type/level 做层次化样式，方便人工检查。
  - 文档末尾增加一张结构树图片，节点标注为 unit_id。

配置：
  - 输入结构树根目录：D:\code\Github\SLAC-test\data\A_structure\papers_structure
  - 输出根目录：D:\code\Github\SLAC-test\data\A_structure\papers_docx_and_structure_graphics
  - 输出目录结构与输入目录一致（保持相对路径）
  - 日志文件：D:\code\Github\SLAC-test\log\docx_and_graphics.log
"""

import json
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Cm

# ===== 可选：树图绘制依赖（Pillow） =====
try:
    from PIL import Image, ImageDraw, ImageFont
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False


# ===== 路径配置 =====

INPUT_ROOT = Path(r"D:\code\Github\SLAC-test\data\A_structure\national_standards_structure\en\part01.tree.json") #D:\code\Github\SLAC-test\data\A_structure\papers_structure
OUTPUT_ROOT = Path(r"D:\code\Github\SLAC-test\data\A_structure\national_standards_structure\en") #D:\code\Github\SLAC-test\data\A_structure\papers_docx_and_structure_graphics
LOG_FILE = Path(r"D:\code\Github\SLAC-test\log\docx_and_graphics.log")


# ===== 日志工具 =====

def log_line(msg: str) -> None:
    """打印并写入日志文件。"""
    print(msg)
    LOG_FILE.parent.mkdir(parents=True, exist_ok=True)
    with LOG_FILE.open("a", encoding="utf-8") as f:
        f.write(msg + "\n")


# ===== 样式设置 =====

def style_heading_paragraph(p, unit_type: str, level: int):
    """
    根据 type 和 level 设置段落的视觉样式。
    - heading: 使用 Word 自带的 Heading 1/2/3 样式，并按 level 缩进。
    - paragraph: 普通正文，按 level 缩进。
    """
    fmt = p.paragraph_format

    if unit_type == "heading":
        # level=0 用 Title，1/2/3 用 Heading 1/2/3，其余都用 Heading 3
        if level <= 0:
            style_name = "Title"
        else:
            mapped_level = min(level, 3)
            style_name = f"Heading {mapped_level}"

        try:
            p.style = style_name
        except Exception:
            # 如果对应样式不存在，就退回 Normal + 手动加粗
            p.style = "Normal"
            for run in p.runs:
                run.bold = True

        # 缩进：每多一层，多 0.5cm
        indent_cm = max(level, 0) * 0.5
        fmt.left_indent = Cm(indent_cm)

    else:
        # 正文段落
        p.style = "Normal"
        # 按 level 缩进一点，和 heading 对齐
        indent_cm = max(level, 0) * 0.5
        fmt.left_indent = Cm(indent_cm)
        fmt.space_after = Pt(4)


def add_unit_to_doc(doc: Document, unit: dict):
    """
    在 Word 文档中追加一个 unit：
    - meta 行
    - 内容行
    """
    unit_id = unit.get("unit_id")
    text = unit.get("text", "")
    unit_type = unit.get("type", "paragraph")
    level = int(unit.get("level", 0))
    parent_id = unit.get("parent_id")

    # 1) Meta 信息行
    meta_text = f"[unit_id={unit_id}, type={unit_type}, level={level}, parent_id={parent_id}]"

    meta_p = doc.add_paragraph()
    meta_run = meta_p.add_run(meta_text)

    # 小号、灰色、斜体
    meta_run.font.size = Pt(8)
    meta_run.font.italic = True
    meta_run.font.color.rgb = RGBColor(128, 128, 128)
    meta_p.paragraph_format.space_after = Pt(0)

    # 2) 内容行
    content_p = doc.add_paragraph(text if text else "(空)")
    style_heading_paragraph(content_p, unit_type, level)

    # 稍微拉开一点距离，便于视觉分块
    content_p.paragraph_format.space_after = Pt(8)


# ===== 结构树图片生成 =====

def generate_tree_image(units, out_dir: Path, doc_id: str) -> Optional[Path]:
    """
    根据 units 生成一张简单的结构树图片：
    - 使用 parent_id 推出“树深度 depth”，而不是直接用 level
    - depth=0 的根节点在最上面，子节点 depth = parent_depth + 1，
      因此保证父子节点不在同一层，子节点必在父节点下方
    - 节点按 depth 分层，横向均匀排布
    - 每个节点画一个圆，内部写 unit_id
    - 父子节点之间画线
    返回图片路径；如果 Pillow 不可用，返回 None。
    """
    if not PIL_AVAILABLE:
        log_line("[WARN] Pillow 未安装，跳过结构树图片生成。可以运行 'pip install pillow' 启用该功能。")
        return None

    # 收集所有有 unit_id 的节点
    nodes: dict[int, dict] = {}
    for u in units:
        uid = u.get("unit_id")
        if uid is None:
            continue
        nodes[uid] = u

    if not nodes:
        log_line("[WARN] 无有效节点，跳过结构树图片生成。")
        return None

    # ===== 1. 用 parent_id 递归计算“树深度 depth” =====
    depths: dict[int, int] = {}

    def get_depth(uid: int, visiting: Optional[set[int]] = None) -> int:
        """
        递归计算某个节点的 depth：
        - 没有 parent 或 parent 不在 nodes 中：depth = 0
        - 否则 depth = parent_depth + 1
        做一个简单的 cycle 保护（如果有环，就当作根）
        """
        if uid in depths:
            return depths[uid]

        if visiting is None:
            visiting = set()
        if uid in visiting:
            # 理论上不会有环，这里做个保险：出现环就把该节点当作根
            d = 0
        else:
            visiting.add(uid)
            u = nodes[uid]
            parent_id = u.get("parent_id")
            if parent_id is None or parent_id not in nodes:
                d = 0
            else:
                d = get_depth(parent_id, visiting) + 1
            visiting.remove(uid)

        depths[uid] = d
        return d

    max_depth = 0
    for uid in nodes.keys():
        d = get_depth(uid)
        if d > max_depth:
            max_depth = d

    # ===== 2. 按 depth 分组（同一层的节点放在一起） =====
    depth_nodes: dict[int, list[int]] = {}
    for uid, d in depths.items():
        depth_nodes.setdefault(d, []).append(uid)

    # 同层按 unit_id 排序，保证稳定
    for d in depth_nodes:
        depth_nodes[d].sort()

    # ===== 3. 画布参数（加大纵向间距，让树不扁） =====
    x_step = 130    # 同层节点水平间距（像素）
    y_step = 220    # 不同层之间垂直间距（像素）
    margin = 80     # 画布边缘留白（像素）

    max_count = max((len(v) for v in depth_nodes.values() if v), default=1)
    width = int(margin * 2 + max_count * x_step)
    height = int(margin * 2 + (max_depth + 1) * y_step)

    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font = ImageFont.load_default()

    # 统一封装一个计算文字尺寸的函数，兼容不同 Pillow 版本
    def get_text_size(text: str):
        if hasattr(draw, "textbbox"):
            left, top, right, bottom = draw.textbbox((0, 0), text, font=font)
            return right - left, bottom - top
        if hasattr(draw, "textsize"):
            return draw.textsize(text, font=font)
        if hasattr(font, "getsize"):
            return font.getsize(text)
        return (20, 10)

    # ===== 4. 计算每个节点的坐标（基于 depth） =====
    positions: dict[int, tuple[int, int]] = {}
    for d in range(max_depth + 1):
        nodes_at_depth = depth_nodes.get(d, [])
        n = len(nodes_at_depth)
        if n == 0:
            continue

        total_width = (n - 1) * x_step
        start_x = margin + (max_count * x_step - total_width) / 2
        y = margin + d * y_step

        for idx, uid in enumerate(nodes_at_depth):
            x = int(start_x + idx * x_step)
            positions[uid] = (x, y)

    # ===== 5. 先画边（父子连线） =====
    for u in units:
        uid = u.get("unit_id")
        parent_id = u.get("parent_id")
        if uid is None or parent_id is None:
            continue
        if uid not in positions or parent_id not in positions:
            continue
        x1, y1 = positions[parent_id]
        x2, y2 = positions[uid]
        draw.line([(x1, y1 + 15), (x2, y2 - 15)], fill="black", width=1)

    # ===== 6. 再画节点（圆 + unit_id） =====
    radius = 15
    for uid, (x, y) in positions.items():
        # 圆
        draw.ellipse(
            (x - radius, y - radius, x + radius, y + radius),
            outline="black",
            fill="white",
            width=1,
        )
        # 文本（unit_id）
        text = str(uid)
        tw, th = get_text_size(text)
        draw.text((x - tw / 2, y - th / 2), text, fill="black", font=font)

    # 标题文字
    title = f"Structure Tree for {doc_id} (unit_id)"
    draw.text((10, 10), title, fill="black", font=font)

    out_dir.mkdir(parents=True, exist_ok=True)
    img_path = out_dir / f"{doc_id}.tree.png"
    img.save(img_path)
    log_line(f"[OK] Tree image saved: {img_path}")
    return img_path


# ===== 单文件处理 =====

def process_tree_file(tree_path: Path, input_root: Path, output_root: Path) -> None:
    """
    处理单个 *.tree.json 文件，生成对应的 .docx 和树图：
    - 输出目录结构与 input_root 相对路径一致。
    """
    try:
        with tree_path.open("r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception as e:
        log_line(f"[ERROR] 读取 JSON 失败: {tree_path} | {e}")
        return

    doc_id = data.get("doc_id", tree_path.stem)
    units = data.get("units", [])

    # 按 unit_id 排序，保证顺序稳定
    try:
        units = sorted(units, key=lambda u: u.get("unit_id", 0))
    except Exception as e:
        log_line(f"[ERROR] units 排序失败: {tree_path} | {e}")
        return

    # 计算输出目录（保持相对路径）
    rel_path = tree_path.relative_to(input_root)          # e.g. s2orc\arxiv_xxx\foo.tree.json
    out_dir = output_root / rel_path.parent
    out_dir.mkdir(parents=True, exist_ok=True)

    # 创建 Word 文档
    doc = Document()

    # 文档标题
    title_p = doc.add_heading(level=0)
    title_run = title_p.add_run(f"Document: {doc_id}")
    title_run.font.size = Pt(16)

    # 简短说明
    intro_p = doc.add_paragraph(
        "每个单元由一行 meta 信息 + 一行内容组成，"
        "meta 显示 unit_id / type / level / parent_id，"
        "内容按层级缩进并使用不同 Heading 样式，便于人工检查结构。"
    )
    intro_p.style = "Normal"
    intro_p.paragraph_format.space_after = Pt(12)

    # 逐个 unit 输出
    for unit in units:
        add_unit_to_doc(doc, unit)

    # 生成结构树图片并插入文档底部
    try:
        img_path = generate_tree_image(units, out_dir, doc_id)
    except Exception as e:
        img_path = None
        log_line(f"[ERROR] 生成结构树图片失败: {tree_path} | {e}")

    if img_path is not None and img_path.exists():
        doc.add_page_break()
        caption_p = doc.add_paragraph("结构树总览（节点标签为 unit_id）：")
        caption_p.style = "Normal"
        caption_p.paragraph_format.space_after = Pt(6)

        # 控制宽度，大约 16cm，保证基本占满一行
        try:
            doc.add_picture(str(img_path), width=Cm(16))
        except Exception as e:
            log_line(f"[ERROR] 向 Word 插入图片失败: {tree_path} | {e}")

    # 保存文件
    out_docx = out_dir / (tree_path.stem + ".labeled.docx")
    try:
        doc.save(out_docx)
        log_line(f"[OK] {tree_path} -> {out_docx}")
    except Exception as e:
        log_line(f"[ERROR] 保存 Word 失败: {tree_path} | {e}")


# ===== 主入口 =====

def main():
    if not INPUT_ROOT.exists():
        log_line(f"[FATAL] INPUT_ROOT not found: {INPUT_ROOT}")
        return

    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)

    # 根据 INPUT_ROOT 类型（文件 / 文件夹）决定处理方式
    if INPUT_ROOT.is_file():
        # 单个 json 文件
        tree_files = [INPUT_ROOT]
        input_root = INPUT_ROOT.parent
        log_line(f"[INFO] INPUT_ROOT 是单个文件，将只处理：{INPUT_ROOT}")
    else:
        # 目录，递归处理目录下所有 json
        input_root = INPUT_ROOT
        tree_files = sorted(INPUT_ROOT.rglob("*.json"))
        log_line(f"[INFO] Found {len(tree_files)} tree files under {INPUT_ROOT}")

    if not tree_files:
        log_line(f"[WARN] No *.json files under {INPUT_ROOT}")
        return

    processed = 0
    for tree_path in tree_files:
        try:
            process_tree_file(tree_path, input_root, OUTPUT_ROOT)
            processed += 1
            if processed % 50 == 0:
                log_line(f"[INFO] 已处理 {processed} / {len(tree_files)} 个文件")
        except Exception as e:
            # 兜底保护，避免单个文件问题中断整个批次
            log_line(f"[ERROR] Unexpected error when processing {tree_path}: {e}")

    log_line(f"[DONE] 全部处理完成，总计 {processed} 个文件。输出根目录：{OUTPUT_ROOT}")



if __name__ == "__main__":
    main()
