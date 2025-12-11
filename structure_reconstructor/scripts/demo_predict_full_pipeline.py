# scripts/demo_predict_full_pipeline.py
"""
硬编码 Demo：
  - 读取一个 txt 或 pdf 文件
  - 用训练好的结构重建器预测结构树（JSON）
  - 基于结构树生成：Word 文档 + 结构树 PNG 图

依赖：
  pip install python-docx networkx matplotlib PyPDF2 transformers
"""

import json
import sys
from pathlib import Path
from typing import List, Dict, Optional

import torch
from docx import Document
from docx.shared import Pt, Cm
import matplotlib.pyplot as plt
import networkx as nx
from PyPDF2 import PdfReader

# ========= 硬编码路径 =========

# 项目根目录（structure_reconstructor）
ROOT = Path(r"D:\code\Github\SLAC-test\structure_reconstructor")

# 训练好的模型 run 目录（包含 best.pt / config_structrec.json）
# 换成你实际的 run 名称，例如 "debug_run" / "slac_struct_v1" 等
RUN_DIR = ROOT / "runs" / "debug_run"

# 要预测的原始文档（txt 或 pdf）——自己改成你要测的那篇
INPUT_PATH = Path(r"D:\code\Github\SLAC-test\data\A_structure\laws\test.txt")
# 也可以用 pdf：
# INPUT_PATH = Path(r"D:\code\Github\SLAC-test\demo_inputs\demo.pdf")

# 输出路径
OUT_DIR = Path(r"D:\code\Github\SLAC-test\demo_outputs")
OUT_JSON = OUT_DIR / "demo.pred.tree.json"
OUT_DOCX = OUT_DIR / "demo.pred.docx"
OUT_PNG = OUT_DIR / "demo.pred.png"

# ======================================

# 确保可以 import src.*
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from src.models.struct_reconstructor import StructRecConfig, StructureReconstructor
from src.data.structure_dataset import build_tokenizer


def read_text(path: Path) -> str:
    """读取 txt 或 pdf 文件为纯文本"""
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return path.read_text(encoding="utf-8")
    elif suffix == ".pdf":
        reader = PdfReader(str(path))
        texts = []
        for page in reader.pages:
            t = page.extract_text() or ""
            texts.append(t)
        return "\n".join(texts)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def simple_split_into_units(text: str) -> List[str]:
    """
    非常简单的分段策略：
      - 按空行切分
      - 去掉首尾空白 & 空段落
    仅用于 Demo，真正线上流程应使用 SLAC 的上游结构抽取器。
    """
    text = text.replace("\r\n", "\n")
    blocks = [b.strip() for b in text.split("\n\n")]
    blocks = [b for b in blocks if b]
    return blocks


def load_model_and_tokenizer(run_dir: Path, device: torch.device):
    cfg_path = run_dir / "config_structrec.json"
    ckpt_path = run_dir / "best.pt"

    if not cfg_path.exists():
        raise FileNotFoundError(f"Config not found: {cfg_path}")
    if not ckpt_path.exists():
        raise FileNotFoundError(f"Checkpoint not found: {ckpt_path}")

    cfg_dict = json.loads(cfg_path.read_text(encoding="utf-8"))
    config = StructRecConfig(**cfg_dict)

    model = StructureReconstructor(config)
    state = torch.load(str(ckpt_path), map_location=device)
    model.load_state_dict(state)
    model.to(device)
    model.eval()

    tokenizer = build_tokenizer(config)
    return model, tokenizer, config


def predict_structure_for_units(
    model: StructureReconstructor,
    tokenizer,
    unit_texts: List[str],
    device: torch.device,
    max_unit_len: int = 256,
) -> List[Dict]:
    """
    给定若干 unit 文本，调用模型预测结构树，返回按 unit_id 排序的节点列表。
    """
    unit_ids = list(range(len(unit_texts)))

    enc = tokenizer(
        unit_texts,
        padding=True,
        truncation=True,
        max_length=max_unit_len,
        return_tensors="pt",
    )
    input_ids = enc["input_ids"].to(device)
    attention_mask = enc["attention_mask"].to(device)

    with torch.no_grad():
        pred_units = model.predict_structure(
            input_ids=input_ids,
            attention_mask=attention_mask,
            unit_texts=unit_texts,
            unit_ids=unit_ids,
        )
    return pred_units


def save_tree_json(
    doc_id: str,
    units: List[Dict],
    out_path: Path,
    source_type: str = "demo",
):
    """
    保存预测结构树为 JSON，结构与训练集一致：
    {
      "doc_id": ...,
      "source_type": "A/B/demo",
      "units": [ {unit_id, text, type, level, parent_id}, ... ]
    }
    """
    out_doc = {
        "doc_id": doc_id,
        "source_type": source_type,
        "units": units,
    }
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(out_doc, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"[OK] Saved predicted tree JSON to: {out_path}")


def save_tree_docx(units: List[Dict], out_path: Path, title: Optional[str] = None):
    """
    根据结构树生成 Word 文档：
      - 每个 unit 两行：meta 行 + 文本行
      - 文本行按 level 缩进；title/heading 用粗体。
    """
    doc = Document()

    if title:
        doc.add_heading(title, level=0)

    for u in units:
        unit_id = u["unit_id"]
        t = u["text"]
        typ = u["type"]
        level = int(u["level"])
        parent_id = u["parent_id"]

        # meta 信息行
        meta = f"[unit_id={unit_id}, type={typ}, level={level}, parent_id={parent_id}]"
        p_meta = doc.add_paragraph(meta)
        r_meta = p_meta.runs[0]
        r_meta.font.size = Pt(9)

        # 文本行
        p_text = doc.add_paragraph(t)
        # 按 level 缩进
        try:
            indent_cm = 0.6 * max(level, 0)
            p_text.paragraph_format.left_indent = Cm(indent_cm)
        except Exception:
            pass

        # type / level 简单控制一下粗体
        for r in p_text.runs:
            if typ in ("title", "heading"):
                r.font.bold = True

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[OK] Saved predicted tree DOCX to: {out_path}")


def hierarchy_pos(G, root=None, width=1.0, vert_gap=0.2, vert_loc=0, xcenter=0.5):
    """
    为树形图生成分层布局的辅助函数。
    来自常见的 networkx 树布局实现。
    """
    if not nx.is_tree(G):
        # 如果有多个 root，就简单用 spring_layout
        return nx.spring_layout(G)

    if root is None:
        # 任选一个入度为 0 的节点作为 root
        roots = [n for n, d in G.in_degree() if d == 0]
        root = roots[0] if roots else list(G.nodes())[0]

    def _hierarchy_pos(G, root, width=1.0, vert_gap=0.2, vert_loc=0, xcenter=0.5, pos=None, parent=None):
        if pos is None:
            pos = {root: (xcenter, vert_loc)}
        else:
            pos[root] = (xcenter, vert_loc)
        children = list(G.successors(root))
        if not children:
            return pos
        dx = width / len(children)
        nextx = xcenter - width / 2 - dx / 2
        for child in children:
            nextx += dx
            pos = _hierarchy_pos(
                G,
                child,
                width=dx,
                vert_gap=vert_gap,
                vert_loc=vert_loc - vert_gap,
                xcenter=nextx,
                pos=pos,
                parent=root,
            )
        return pos

    return _hierarchy_pos(G, root, width, vert_gap, vert_loc, xcenter)


def save_tree_png(units: List[Dict], out_path: Path):
    """
    画出结构树，节点标签为 unit_id。
    """
    G = nx.DiGraph()
    for u in units:
        uid = u["unit_id"]
        G.add_node(uid)
    for u in units:
        uid = u["unit_id"]
        pid = u["parent_id"]
        if pid is not None:
            G.add_edge(pid, uid)

    if len(G.nodes) == 0:
        print("[WARN] Empty tree, skip PNG.")
        return

    pos = hierarchy_pos(G)

    plt.figure(figsize=(max(8, len(G.nodes) * 0.3), 6))
    nx.draw(
        G,
        pos,
        with_labels=True,
        labels={n: str(n) for n in G.nodes},
        arrows=True,
        node_size=600,
        font_size=8,
    )
    out_path.parent.mkdir(parents=True, exist_ok=True)
    plt.tight_layout()
    plt.savefig(str(out_path), dpi=200)
    plt.close()
    print(f"[OK] Saved structure tree PNG to: {out_path}")


def main():
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    print(f"[INFO] Using device: {device}")

    raw_text = read_text(INPUT_PATH)
    print(f"[INFO] Loaded raw text from {INPUT_PATH}, length={len(raw_text)} chars")

    unit_texts = simple_split_into_units(raw_text)
    print(f"[INFO] Split into {len(unit_texts)} units")

    model, tokenizer, config = load_model_and_tokenizer(RUN_DIR, device)

    pred_units = predict_structure_for_units(
        model=model,
        tokenizer=tokenizer,
        unit_texts=unit_texts,
        device=device,
        max_unit_len=256,
    )

    doc_id = INPUT_PATH.stem

    save_tree_json(doc_id=doc_id, units=pred_units, out_path=OUT_JSON, source_type="demo")
    save_tree_docx(pred_units, out_path=OUT_DOCX, title=f"Predicted structure: {doc_id}")
    save_tree_png(pred_units, out_path=OUT_PNG)


if __name__ == "__main__":
    main()
